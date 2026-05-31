from __future__ import annotations

import csv
import io
import re
from pathlib import Path
from typing import Any
from uuid import uuid4

import httpx
from fastapi import UploadFile
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.db.models import App, KnowledgeBase, KnowledgeChunk, KnowledgeDocument
from app.db.session import SessionLocal
from app.schemas import KnowledgeBaseCreate, KnowledgeBaseUpdate
from app.services.qdrant_service import (
    delete_collection,
    delete_knowledge_points,
    ensure_collection,
    update_knowledge_sparse_vectors,
    upsert_knowledge_chunks,
)
from app.services.retrieval_defaults import (
    DEFAULT_CHUNK_OVERLAP,
    DEFAULT_CHUNK_SIZE,
    DEFAULT_SPARSE_STOPWORDS_ENABLED,
    DEFAULT_SPARSE_TOKENIZER,
    DEFAULT_SPARSE_WEIGHTING,
    DEFAULT_PARENT_CHILD_ENABLED,
    DEFAULT_PARENT_CHUNK_OVERLAP,
    DEFAULT_PARENT_CHUNK_SIZE,
)
from app.services.retrieval_providers import (
    build_embedding_provider,
    normalize_provider,
    validate_embedding_dimensions,
)
from app.services.sparse_bm25 import build_bm25_sparse_vectors


DIRECT_TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".py",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".java",
    ".go",
    ".json",
    ".yaml",
    ".yml",
    ".csv",
    ".html",
    ".css",
}
SUPPORTED_SUFFIXES = DIRECT_TEXT_SUFFIXES | {".pdf", ".docx", ".pptx", ".xlsx", ".xls"}
CODE_SUFFIXES = {".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".go", ".css", ".html"}


def list_knowledge_bases(db: Session, owner_user_id: str) -> list[KnowledgeBase]:
    return list(
        db.scalars(
            select(KnowledgeBase)
            .where(KnowledgeBase.owner_user_id == owner_user_id, KnowledgeBase.scope == "creator")
            .order_by(KnowledgeBase.created_at.desc())
        )
    )


def get_knowledge_base(db: Session, kb_id: str, owner_user_id: str) -> KnowledgeBase | None:
    return db.scalar(
        select(KnowledgeBase).where(
            KnowledgeBase.id == kb_id,
            KnowledgeBase.owner_user_id == owner_user_id,
            KnowledgeBase.scope == "creator",
        )
    )


def create_knowledge_base(db: Session, payload: KnowledgeBaseCreate, owner_user_id: str) -> KnowledgeBase:
    embedding = _current_embedding_snapshot(require_api_key=True)
    kb = KnowledgeBase(
        owner_user_id=owner_user_id,
        scope="creator",
        app_id="",
        conversation_id="",
        name=payload.name.strip(),
        description=payload.description.strip(),
        embedding_provider=embedding["provider"],
        embedding_model=embedding["model"],
        embedding_dimension=embedding["dimension"],
        embedding_credential_id="",
        embedding_base_url=embedding["base_url"],
        qdrant_collection=f"kb_{uuid4().hex}",
        chunk_size=DEFAULT_CHUNK_SIZE,
        chunk_overlap=DEFAULT_CHUNK_OVERLAP,
        chunk_strategy="auto",
        enable_parent_child=DEFAULT_PARENT_CHILD_ENABLED,
        config_json={
            "embedding_snapshot": embedding,
            "parent_chunk_size": DEFAULT_PARENT_CHUNK_SIZE,
            "parent_chunk_overlap": DEFAULT_PARENT_CHUNK_OVERLAP,
            "qdrant_sparse_enabled": True,
            "sparse_weighting": DEFAULT_SPARSE_WEIGHTING,
            "sparse_tokenizer": DEFAULT_SPARSE_TOKENIZER,
            "sparse_stopwords_enabled": DEFAULT_SPARSE_STOPWORDS_ENABLED,
        },
    )
    db.add(kb)
    db.flush()
    ensure_collection(kb)
    db.commit()
    db.refresh(kb)
    return kb


def update_knowledge_base(db: Session, kb: KnowledgeBase, payload: KnowledgeBaseUpdate) -> KnowledgeBase:
    values = payload.model_dump(exclude_unset=True)
    if "name" in values and values["name"] is not None:
        kb.name = str(values["name"]).strip()
    if "description" in values and values["description"] is not None:
        kb.description = str(values["description"]).strip()
    db.commit()
    db.refresh(kb)
    return kb


def delete_knowledge_base(db: Session, kb: KnowledgeBase) -> None:
    try:
        delete_collection(kb)
    except Exception:
        pass
    _remove_knowledge_base_from_app_workflows(db, kb)
    db.delete(kb)
    db.commit()


def list_knowledge_documents(db: Session, kb: KnowledgeBase) -> list[KnowledgeDocument]:
    return list(
        db.scalars(
            select(KnowledgeDocument)
            .where(KnowledgeDocument.knowledge_base_id == kb.id)
            .order_by(KnowledgeDocument.created_at.desc())
        )
    )


async def save_knowledge_document(db: Session, kb: KnowledgeBase, file: UploadFile) -> KnowledgeDocument:
    _assert_embedding_config_not_drifted(kb)
    filename = Path(file.filename or "").name
    suffix = Path(filename).suffix.lower()
    if not filename:
        raise ValueError("Filename is required.")
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError("Only txt/md/code/pdf/docx/pptx/xlsx files are supported by knowledge databases.")

    kb_dir = Path(get_settings().storage_dir) / "knowledge_bases" / kb.id
    kb_dir.mkdir(parents=True, exist_ok=True)
    raw = await file.read()
    file_path = kb_dir / f"{uuid4()}_{filename}"
    file_path.write_bytes(raw)

    document = KnowledgeDocument(
        knowledge_base_id=kb.id,
        filename=filename,
        file_path=str(file_path),
        mime_type=file.content_type or "",
        status="queued",
        metadata_json={"suffix": suffix},
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    enqueue_knowledge_document(document.id)
    db.refresh(document)
    return document


def delete_knowledge_document(db: Session, kb: KnowledgeBase, document_id: str) -> bool:
    document = db.get(KnowledgeDocument, document_id)
    if not document or document.knowledge_base_id != kb.id:
        return False
    point_ids = [
        chunk.qdrant_point_id or chunk.id
        for chunk in db.scalars(select(KnowledgeChunk).where(KnowledgeChunk.document_id == document.id))
    ]
    try:
        delete_knowledge_points(kb, point_ids)
    except Exception:
        pass
    db.delete(document)
    db.commit()
    _refresh_knowledge_sparse_vectors(db, kb)
    return True


def rebuild_knowledge_base(db: Session, kb: KnowledgeBase) -> KnowledgeBase:
    _assert_embedding_config_not_drifted(kb)
    try:
        delete_collection(kb)
    except Exception:
        pass
    db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.knowledge_base_id == kb.id))
    kb.locked = False
    ensure_collection(kb)
    for document in db.scalars(select(KnowledgeDocument).where(KnowledgeDocument.knowledge_base_id == kb.id)):
        document.status = "queued"
        document.error = ""
        document.metadata_json = {**dict(document.metadata_json or {}), "rebuilt": True}
    db.commit()
    for document in list_knowledge_documents(db, kb):
        enqueue_knowledge_document(document.id)
    db.refresh(kb)
    return kb


def enqueue_knowledge_document(document_id: str) -> None:
    if get_settings().document_inline_ingest:
        process_knowledge_document_sync(document_id)
        return
    from app.worker.celery_app import process_knowledge_document

    process_knowledge_document.delay(document_id)


def process_knowledge_document_sync(document_id: str) -> None:
    db = SessionLocal()
    try:
        _process_knowledge_document(db, document_id)
    finally:
        db.close()


def _process_knowledge_document(db: Session, document_id: str) -> None:
    document = db.get(KnowledgeDocument, document_id)
    if not document:
        return
    kb = document.knowledge_base
    try:
        _assert_embedding_config_not_drifted(kb)
        _update_document_status(db, document, "parsing")
        parse_result = _parse_document(Path(document.file_path), document.filename)
        elements = parse_result["elements"]
        document.metadata_json = {
            **dict(document.metadata_json or {}),
            "parser": parse_result.get("parser", {}),
            "warnings": parse_result.get("warnings", []),
            "element_count": len(elements),
        }
        db.commit()

        _update_document_status(db, document, "chunking")
        chunks = _store_chunks(db, kb, document, elements)

        _update_document_status(db, document, "embedding")
        provider = build_embedding_provider(kb)
        vectors = provider.embed_documents([chunk.content for chunk in chunks])
        validate_embedding_dimensions(vectors, kb.embedding_dimension)
        for chunk in chunks:
            chunk.qdrant_point_id = chunk.id
        db.flush()
        ensure_collection(kb)
        sparse_vectors_by_chunk_id = _build_sparse_vectors_by_chunk_id(db, kb)
        upsert_knowledge_chunks(
            kb,
            chunks,
            vectors,
            [sparse_vectors_by_chunk_id.get(chunk.id, {"indices": [], "values": []}) for chunk in chunks],
        )
        _refresh_knowledge_sparse_vectors(
            db,
            kb,
            skip_chunk_ids={chunk.id for chunk in chunks},
            sparse_vectors_by_chunk_id=sparse_vectors_by_chunk_id,
        )

        document.status = "ready"
        document.error = ""
        kb.locked = True
        db.commit()
    except Exception as exc:
        document.status = "failed"
        document.error = str(exc)
        db.commit()


def _build_sparse_vectors_by_chunk_id(
    db: Session,
    kb: KnowledgeBase,
) -> dict[str, dict[str, list[int] | list[float]]]:
    chunks = _list_sparse_index_chunks(db, kb)
    sparse_vectors = build_bm25_sparse_vectors([chunk.content for chunk in chunks])
    return {chunk.id: sparse_vector for chunk, sparse_vector in zip(chunks, sparse_vectors)}


def _refresh_knowledge_sparse_vectors(
    db: Session,
    kb: KnowledgeBase,
    skip_chunk_ids: set[str] | None = None,
    sparse_vectors_by_chunk_id: dict[str, dict[str, list[int] | list[float]]] | None = None,
) -> None:
    chunks = _list_sparse_index_chunks(db, kb)
    if not chunks:
        return
    skipped = skip_chunk_ids or set()
    if sparse_vectors_by_chunk_id is None:
        sparse_vectors_by_chunk_id = _build_sparse_vectors_by_chunk_id(db, kb)
    refresh_chunks = [chunk for chunk in chunks if chunk.id not in skipped]
    if not refresh_chunks:
        return
    update_knowledge_sparse_vectors(
        kb,
        refresh_chunks,
        [sparse_vectors_by_chunk_id.get(chunk.id, {"indices": [], "values": []}) for chunk in refresh_chunks],
    )


def _remove_knowledge_base_from_app_workflows(db: Session, kb: KnowledgeBase) -> None:
    apps = list(db.scalars(select(App).where(App.owner_user_id == kb.owner_user_id)))
    for app in apps:
        spec = dict(app.workflow_spec or {})
        nodes = spec.get("nodes")
        if not isinstance(nodes, list):
            continue
        changed = False
        next_nodes: list[Any] = []
        for node in nodes:
            if not isinstance(node, dict):
                next_nodes.append(node)
                continue
            next_node = dict(node)
            ids = next_node.get("knowledge_base_ids")
            if isinstance(ids, list):
                next_ids = [str(item) for item in ids if str(item) != kb.id]
                if next_ids != ids:
                    next_node["knowledge_base_ids"] = next_ids
                    changed = True
            if str(next_node.get("kb_id") or "") == kb.id:
                next_node.pop("kb_id", None)
                changed = True
            next_nodes.append(next_node)
        if changed:
            app.workflow_spec = {**spec, "nodes": next_nodes}


def _list_sparse_index_chunks(db: Session, kb: KnowledgeBase) -> list[KnowledgeChunk]:
    chunks = list(
        db.scalars(
            select(KnowledgeChunk)
            .where(KnowledgeChunk.knowledge_base_id == kb.id)
            .order_by(KnowledgeChunk.chunk_index.asc())
        )
    )
    return [chunk for chunk in chunks if _is_sparse_indexable_chunk(chunk)]


def _is_sparse_indexable_chunk(chunk: KnowledgeChunk) -> bool:
    metadata = dict(chunk.metadata_json or {})
    return (
        metadata.get("chunk_type") != "parent"
        and bool(str(chunk.content or "").strip())
        and bool(chunk.qdrant_point_id)
    )


def _update_document_status(db: Session, document: KnowledgeDocument, status: str) -> None:
    document.status = status
    document.error = ""
    db.commit()
    db.refresh(document)


def _current_embedding_snapshot(require_api_key: bool = False) -> dict[str, Any]:
    settings = get_settings()
    provider = normalize_provider(settings.knowledge_embedding_provider)
    model = settings.knowledge_embedding_model.strip()
    dimension = int(settings.knowledge_embedding_dimension or 0)
    base_url = _resolve_embedding_base_url(provider, settings.knowledge_embedding_base_url)
    if not provider:
        raise ValueError("KNOWLEDGE_EMBEDDING_PROVIDER is required.")
    if not model:
        raise ValueError("KNOWLEDGE_EMBEDDING_MODEL is required.")
    if dimension <= 0:
        raise ValueError("KNOWLEDGE_EMBEDDING_DIMENSION is required.")
    if require_api_key and not settings.knowledge_embedding_api_key:
        raise ValueError("KNOWLEDGE_EMBEDDING_API_KEY is required.")
    return {
        "provider": provider,
        "model": model,
        "dimension": dimension,
        "base_url": base_url.rstrip("/"),
    }


def _embedding_snapshot_from_kb(kb: KnowledgeBase) -> dict[str, Any]:
    config = dict(kb.config_json or {})
    snapshot = config.get("embedding_snapshot")
    if isinstance(snapshot, dict):
        return {
            "provider": normalize_provider(str(snapshot.get("provider") or "")),
            "model": str(snapshot.get("model") or "").strip(),
            "dimension": _to_int(snapshot.get("dimension"), 0),
            "base_url": str(snapshot.get("base_url") or "").strip().rstrip("/"),
        }
    return {
        "provider": normalize_provider(kb.embedding_provider),
        "model": str(kb.embedding_model or "").strip(),
        "dimension": int(kb.embedding_dimension or 0),
        "base_url": str(kb.embedding_base_url or "").strip().rstrip("/"),
    }


def _assert_embedding_config_not_drifted(kb: KnowledgeBase) -> None:
    current = _current_embedding_snapshot()
    locked = _embedding_snapshot_from_kb(kb)
    if current == locked:
        return
    raise ValueError(
        "Knowledge embedding config drift detected. This knowledge database was embedded with "
        f"{_format_embedding_snapshot(locked)}, but the backend currently uses {_format_embedding_snapshot(current)}. "
        "Create or rebuild a knowledge database with the current backend embedding config."
    )


def _format_embedding_snapshot(snapshot: dict[str, Any]) -> str:
    return (
        f"provider={snapshot.get('provider') or 'empty'}, "
        f"model={snapshot.get('model') or 'empty'}, "
        f"dimension={snapshot.get('dimension') or 0}, "
        f"base_url={snapshot.get('base_url') or 'empty'}"
    )


def _resolve_embedding_base_url(provider_id: str, base_url: str) -> str:
    explicit = str(base_url or "").strip()
    if explicit:
        return explicit
    if provider_id == "openai":
        return "https://api.openai.com/v1"
    if provider_id in {"dashscope", "qwen"}:
        return "https://dashscope.aliyuncs.com/compatible-mode/v1"
    if provider_id in {"zhipu", "zhipuai"}:
        return "https://open.bigmodel.cn/api/paas/v4"
    return ""


def _parse_document(path: Path, filename: str) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix in DIRECT_TEXT_SUFFIXES:
        return _parse_result([_element_from_direct_file(path, suffix)], route="direct", parser="direct")
    if suffix == ".pdf":
        pdfminer_elements = _parse_pdf_with_pdfminer(path)
        if _total_element_text_length(pdfminer_elements) >= 64:
            external_elements = _try_parse_with_unstructured(path)
            elements = external_elements or pdfminer_elements
            return _parse_result(
                elements,
                route="pdf_text_layer_unstructured" if external_elements else "pdf_text_layer_pdfminer",
                parser="unstructured" if external_elements else "pdfminer",
            )
        miner_elements = _try_parse_with_miner(path)
        if miner_elements:
            return _parse_result(miner_elements, route="pdf_scanned_miner", parser="miner")
        raise ValueError(
            "No text layer was found in this PDF and the configured scanned-PDF parser returned no text."
        )
    if suffix == ".docx":
        external_elements = _try_parse_with_unstructured(path)
        if external_elements:
            return _parse_result(external_elements, route="docx_unstructured", parser="unstructured")
        return _parse_result(_parse_docx_locally(path), route="docx_python_docx", parser="python-docx")
    if suffix in {".pptx", ".xlsx", ".xls"}:
        external_elements = _try_parse_with_unstructured(path, raise_on_error=True)
        if external_elements:
            return _parse_result(external_elements, route=f"{suffix.lstrip('.')}_unstructured", parser="unstructured")
        raise ValueError(f"{suffix} files require DOCUMENT_UNSTRUCTURED_API_URL.")
    raise ValueError(f"Unsupported file type for {filename}.")


def _parse_result(elements: list[dict[str, Any]], route: str, parser: str) -> dict[str, Any]:
    normalized = _attach_sections(elements)
    warnings = _element_warnings(normalized)
    return {
        "elements": normalized,
        "parser": {
            "route": route,
            "parser": parser,
        },
        "warnings": warnings,
    }


def _element_from_direct_file(path: Path, suffix: str) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".csv":
        return _new_element(text, element_type="Table", chunk_type="table", metadata={"suffix": suffix})
    return _new_element(
        text,
        element_type="Code" if suffix in CODE_SUFFIXES else "NarrativeText",
        chunk_type="code" if suffix in CODE_SUFFIXES else "text",
        metadata={"suffix": suffix, "language": suffix.lstrip(".")},
    )


def _parse_pdf_with_pdfminer(path: Path) -> list[dict[str, Any]]:
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTTextContainer
    except ImportError as exc:
        raise RuntimeError("pdfminer.six is not installed. Run `pip install -e .` in backend first.") from exc

    elements: list[dict[str, Any]] = []
    for page_number, page_layout in enumerate(extract_pages(str(path)), start=1):
        parts = []
        for item in page_layout:
            if isinstance(item, LTTextContainer):
                text = item.get_text().strip()
                if text:
                    parts.append(text)
        page_text = "\n".join(parts).strip()
        if page_text:
            elements.append(_new_element(page_text, element_type="NarrativeText", chunk_type="text", page_num=page_number))
    return elements


def _parse_docx_locally(path: Path) -> list[dict[str, Any]]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed. Run `pip install -e .` in backend first.") from exc
    doc = DocxDocument(str(path))
    elements: list[dict[str, Any]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = str(getattr(para.style, "name", "") or "")
        element_type = "Title" if style_name.lower().startswith("heading") else "NarrativeText"
        elements.append(_new_element(text, element_type=element_type, chunk_type="text"))
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append("\t".join(cell.text.strip() for cell in row.cells))
        table_text = "\n".join(line for line in rows if line.strip())
        if table_text:
            elements.append(_new_element(table_text, element_type="Table", chunk_type="table"))
    return elements


def _try_parse_with_unstructured(path: Path, raise_on_error: bool = False) -> list[dict[str, Any]]:
    settings = get_settings()
    try:
        if settings.document_unstructured_api_url:
            return _parse_with_unstructured_http(
                path,
                settings.document_unstructured_api_url,
                settings.document_unstructured_api_key,
            )
        try:
            from unstructured.partition.auto import partition
        except ImportError:
            return []
        raw_elements = partition(filename=str(path))
        return [_element_from_unstructured(item) for item in raw_elements if str(item).strip()]
    except Exception as exc:
        if raise_on_error:
            raise ValueError(f"Unstructured parser failed for {path.name}: {exc}") from exc
        return []


def _parse_with_unstructured_http(path: Path, api_url: str, api_key: str = "") -> list[dict[str, Any]]:
    endpoint = api_url.rstrip("/")
    if not endpoint.endswith("/general/v0/general"):
        endpoint = f"{endpoint}/general/v0/general"
    headers = {"Accept": "application/json"}
    if api_key:
        headers["unstructured-api-key"] = api_key
        if "unstructured" not in endpoint.lower():
            headers["Authorization"] = f"Bearer {api_key}"
    with path.open("rb") as handle:
        response = httpx.post(
            endpoint,
            headers=headers,
            files={"files": (path.name, handle, "application/octet-stream")},
            data={"strategy": "auto", "output_format": "application/json"},
            timeout=120.0,
        )
    response.raise_for_status()
    payload = response.json()
    if isinstance(payload, dict) and isinstance(payload.get("elements"), list):
        payload = payload["elements"]
    if not isinstance(payload, list):
        raise ValueError("Unstructured returned an invalid response: expected a list of elements.")
    return [_element_from_unstructured_dict(item) for item in payload if isinstance(item, dict)]


def _try_parse_with_miner(path: Path) -> list[dict[str, Any]]:
    settings = get_settings()
    if not settings.document_mineru_api_url:
        return []
    endpoint = settings.document_mineru_api_url.rstrip("/")
    if "mineru.net" in endpoint.lower() and "/api/v1/agent/parse/file" in endpoint.lower():
        return _parse_with_miner_agent_file(path, endpoint)
    headers = {"Accept": "application/json"}
    if settings.document_mineru_api_key:
        headers["Authorization"] = f"Bearer {settings.document_mineru_api_key}"
    try:
        with path.open("rb") as handle:
            response = httpx.post(
                endpoint,
                headers=headers,
                files={"file": (path.name, handle, "application/pdf")},
                timeout=180.0,
            )
        response.raise_for_status()
        payload = response.json()
    except Exception as exc:
        raise ValueError(f"Scanned-PDF parser failed for {path.name}: {exc}") from exc
    raw_elements = payload.get("elements") if isinstance(payload, dict) else payload
    if not isinstance(raw_elements, list):
        raise ValueError("Scanned-PDF parser returned an invalid response: expected elements list.")
    return [_element_from_unstructured_dict(item, default_parser="miner") for item in raw_elements if isinstance(item, dict)]


def _parse_with_miner_agent_file(path: Path, endpoint: str) -> list[dict[str, Any]]:
    settings = get_settings()
    headers = {"Accept": "application/json"}
    if settings.document_mineru_api_key:
        headers["Authorization"] = f"Bearer {settings.document_mineru_api_key}"
    try:
        create_response = httpx.post(
            endpoint,
            headers=headers,
            json={
                "file_name": path.name,
                "language": "ch",
                "is_ocr": True,
                "enable_formula": True,
                "enable_table": True,
            },
            timeout=30.0,
        )
        create_response.raise_for_status()
        create_payload = create_response.json()
        create_data = create_payload.get("data") if isinstance(create_payload, dict) else None
        if not isinstance(create_data, dict):
            raise ValueError("Scanned-PDF Agent API returned an invalid task response.")

        upload_url = str(create_data.get("file_url") or create_data.get("url") or "")
        task_id = str(create_data.get("task_id") or "")
        if not upload_url or not task_id:
            raise ValueError("Scanned-PDF Agent API response is missing upload url or task_id.")

        with path.open("rb") as handle:
            upload_response = httpx.put(upload_url, content=handle.read(), timeout=180.0)
        upload_response.raise_for_status()

        status_url = endpoint.rsplit("/file", 1)[0] + f"/{task_id}"
        markdown_url = ""
        last_status = ""
        for _ in range(120):
            status_response = httpx.get(status_url, headers=headers, timeout=30.0)
            status_response.raise_for_status()
            status_payload = status_response.json()
            status_data = status_payload.get("data") if isinstance(status_payload, dict) else None
            if not isinstance(status_data, dict):
                raise ValueError("Scanned-PDF Agent API returned an invalid status response.")
            last_status = str(status_data.get("state") or status_data.get("status") or "")
            full_zip_url = str(status_data.get("full_zip_url") or "")
            markdown_url = str(status_data.get("md_url") or status_data.get("markdown_url") or "")
            if not markdown_url and full_zip_url:
                markdown_url = full_zip_url
            if markdown_url:
                break
            if last_status.lower() in {"failed", "error", "done_failed"}:
                raise ValueError(f"Scanned-PDF Agent task failed: {status_data}")
            import time

            time.sleep(2)

        if not markdown_url:
            raise ValueError(f"Scanned-PDF Agent task did not finish in time. Last status: {last_status or 'unknown'}")

        markdown_response = httpx.get(markdown_url, timeout=120.0)
        markdown_response.raise_for_status()
        markdown = markdown_response.text.strip()
        if not markdown:
            raise ValueError("Scanned-PDF Agent returned an empty markdown result.")
        return [
            _new_element(
                markdown,
                element_type="NarrativeText",
                chunk_type="text",
                metadata={"parser": "miner_agent", "task_id": task_id, "markdown_url": markdown_url},
            )
        ]
    except Exception as exc:
        raise ValueError(f"Scanned-PDF Agent parser failed for {path.name}: {exc}") from exc


def _element_from_unstructured(item: Any) -> dict[str, Any]:
    metadata = getattr(item, "metadata", None)
    metadata_dict = metadata.to_dict() if hasattr(metadata, "to_dict") else {}
    category = str(getattr(item, "category", "") or item.__class__.__name__ or "NarrativeText")
    return _new_element(
        str(item),
        element_type=category,
        chunk_type=_chunk_type_for_element(category),
        page_num=metadata_dict.get("page_number"),
        metadata=metadata_dict,
    )


def _element_from_unstructured_dict(item: dict[str, Any], default_parser: str = "unstructured") -> dict[str, Any]:
    metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else {}
    element_type = str(item.get("type") or item.get("category") or item.get("element_type") or "NarrativeText")
    text = str(item.get("text") or item.get("content") or "")
    return _new_element(
        text,
        element_type=element_type,
        chunk_type=_chunk_type_for_element(element_type),
        page_num=item.get("page_number") or metadata.get("page_number") or metadata.get("page_num"),
        metadata={**metadata, "parser": default_parser},
    )


def _new_element(
    text: str,
    element_type: str,
    chunk_type: str,
    section: str | None = None,
    page_num: Any = None,
    metadata: dict[str, Any] | None = None,
) -> dict[str, Any]:
    return {
        "text": str(text or ""),
        "element_type": str(element_type or "NarrativeText"),
        "chunk_type": chunk_type,
        "section": section,
        "page_num": _to_int(page_num, 0) or None,
        "metadata": metadata or {},
    }


def _chunk_type_for_element(element_type: str) -> str:
    value = str(element_type or "").lower()
    if "table" in value:
        return "table"
    if "code" in value:
        return "code"
    if "image" in value or "figure" in value:
        return "image"
    return "text"


def _attach_sections(elements: list[dict[str, Any]]) -> list[dict[str, Any]]:
    current_section = None
    attached: list[dict[str, Any]] = []
    for element in elements:
        element_type = str(element.get("element_type") or "")
        text = str(element.get("text") or "").strip()
        if not text and _chunk_type_for_element(element_type) != "image":
            continue
        if element_type.lower() == "title":
            current_section = text or current_section
            continue
        next_element = dict(element)
        next_element["section"] = next_element.get("section") or current_section
        attached.append(next_element)
    return attached


def _element_warnings(elements: list[dict[str, Any]]) -> list[str]:
    skipped_images = sum(1 for element in elements if element.get("chunk_type") == "image")
    if skipped_images:
        return [f"Skipped {skipped_images} image element(s); image content retrieval is not supported yet."]
    return []


def _total_element_text_length(elements: list[dict[str, Any]]) -> int:
    return sum(len(str(element.get("text") or "").strip()) for element in elements)


def _store_chunks(
    db: Session,
    kb: KnowledgeBase,
    document: KnowledgeDocument,
    elements: list[dict[str, Any]],
) -> list[KnowledgeChunk]:
    db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_id == document.id))
    vector_chunks: list[KnowledgeChunk] = []
    index = 0
    for element in elements:
        text = str(element.get("text") or "").strip()
        if not text or element.get("chunk_type") == "image":
            continue
        text = _normalize_element_text(element, text)
        if kb.enable_parent_child:
            index, next_chunks = _store_parent_child_chunks(db, kb, document, element, text, index)
            vector_chunks.extend(next_chunks)
            continue
        for content in _split_element_text(element, text, kb.chunk_size, kb.chunk_overlap):
            chunk = _new_knowledge_chunk(
                kb=kb,
                document=document,
                index=index,
                content=content,
                metadata=_base_chunk_metadata(kb, document, element, chunk_type=element.get("chunk_type") or "text"),
            )
            db.add(chunk)
            vector_chunks.append(chunk)
            index += 1

    if not vector_chunks:
        raise ValueError("No retrievable text was parsed from this document.")

    db.flush()
    for chunk in db.scalars(select(KnowledgeChunk).where(KnowledgeChunk.document_id == document.id)):
        _fill_chunk_id(chunk)
    db.commit()
    return vector_chunks


def _store_parent_child_chunks(
    db: Session,
    kb: KnowledgeBase,
    document: KnowledgeDocument,
    element: dict[str, Any],
    text: str,
    index: int,
) -> tuple[int, list[KnowledgeChunk]]:
    config = dict(kb.config_json or {})
    parent_chunk_size = _to_int(config.get("parent_chunk_size"), max(kb.chunk_size * 4, kb.chunk_size))
    parent_chunk_overlap = _to_int(config.get("parent_chunk_overlap"), max(kb.chunk_overlap * 2, 0))
    vector_chunks: list[KnowledgeChunk] = []

    for parent_content in _split_element_text(element, text, parent_chunk_size, parent_chunk_overlap):
        parent_chunk = _new_knowledge_chunk(
            kb=kb,
            document=document,
            index=index,
            content=parent_content,
            metadata=_base_chunk_metadata(kb, document, element, chunk_type="parent"),
        )
        db.add(parent_chunk)
        db.flush()
        _fill_chunk_id(parent_chunk)
        index += 1

        source_type = str(element.get("chunk_type") or "text")
        for child_content in _split_element_text(element, parent_content, kb.chunk_size, kb.chunk_overlap):
            metadata = _base_chunk_metadata(kb, document, element, chunk_type="child")
            metadata["parent_id"] = parent_chunk.id
            metadata["source_chunk_type"] = source_type
            child_chunk = _new_knowledge_chunk(
                kb=kb,
                document=document,
                index=index,
                content=child_content,
                metadata=metadata,
            )
            db.add(child_chunk)
            vector_chunks.append(child_chunk)
            index += 1

    return index, vector_chunks


def _new_knowledge_chunk(
    kb: KnowledgeBase,
    document: KnowledgeDocument,
    index: int,
    content: str,
    metadata: dict[str, Any],
) -> KnowledgeChunk:
    return KnowledgeChunk(
        knowledge_base_id=kb.id,
        document_id=document.id,
        chunk_index=index,
        content=content,
        metadata_json=metadata,
    )


def _base_chunk_metadata(
    kb: KnowledgeBase,
    document: KnowledgeDocument,
    element: dict[str, Any],
    chunk_type: str,
) -> dict[str, Any]:
    return {
        "source_file": document.filename,
        "chunk_id": "",
        "parent_id": None,
        "chunk_type": chunk_type,
        "element_type": element.get("element_type"),
        "section": element.get("section"),
        "page_num": element.get("page_num"),
        "kb_id": kb.id,
        "scope": kb.scope,
    }


def _fill_chunk_id(chunk: KnowledgeChunk) -> None:
    metadata = dict(chunk.metadata_json or {})
    metadata["chunk_id"] = chunk.id
    chunk.metadata_json = metadata


def _normalize_element_text(element: dict[str, Any], text: str) -> str:
    if element.get("chunk_type") == "table":
        return _table_to_text(text)
    return text


def _split_element_text(
    element: dict[str, Any],
    text: str,
    chunk_size: int,
    chunk_overlap: int,
) -> list[str]:
    chunk_type = str(element.get("chunk_type") or "text")
    if chunk_type == "table":
        return [text.strip()] if text.strip() else []
    if chunk_type == "code":
        return _split_code_text(text, chunk_size, chunk_overlap)
    return _split_text(text, chunk_size, chunk_overlap)


def _split_code_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    char_budget = max(chunk_size * 4, 1)
    if len(text) <= char_budget:
        return [text.strip()] if text.strip() else []

    lines = text.splitlines()
    boundaries = [0]
    for index, line in enumerate(lines):
        if index and _is_code_boundary(line):
            boundaries.append(index)
    boundaries.append(len(lines))
    if len(boundaries) <= 2:
        return _split_text(text, chunk_size, chunk_overlap)

    blocks = []
    for start, end in zip(boundaries, boundaries[1:]):
        block = "\n".join(lines[start:end]).strip()
        if block:
            blocks.append(block)

    chunks: list[str] = []
    current = ""
    for block in blocks:
        if len(block) > char_budget:
            if current.strip():
                chunks.append(current.strip())
                current = ""
            chunks.extend(_split_text(block, chunk_size, chunk_overlap))
            continue
        candidate = f"{current}\n\n{block}".strip() if current else block
        if len(candidate) <= char_budget:
            current = candidate
        else:
            if current.strip():
                chunks.append(current.strip())
            current = block
    if current.strip():
        chunks.append(current.strip())
    return chunks


def _is_code_boundary(line: str) -> bool:
    return bool(
        re.match(r"^\s*(async\s+)?def\s+\w+\s*\(", line)
        or re.match(r"^\s*class\s+\w+", line)
        or re.match(r"^\s*(export\s+)?(async\s+)?function\s+\w+\s*\(", line)
        or re.match(r"^\s*(export\s+)?class\s+\w+", line)
        or re.match(r"^\s*(const|let|var)\s+\w+\s*=\s*(async\s*)?\(", line)
        or re.match(r"^\s*func\s+\w+\s*\(", line)
        or re.match(r"^\s*(public|private|protected)?\s*(static\s+)?[\w<>\[\]]+\s+\w+\s*\(", line)
    )


def _table_to_text(text: str) -> str:
    rows = _parse_table_rows(text)
    if len(rows) < 2:
        return text.strip()
    headers = [cell.strip() or f"column_{index + 1}" for index, cell in enumerate(rows[0])]
    lines = []
    for row_index, row in enumerate(rows[1:], start=1):
        cells = [cell.strip() for cell in row]
        if not any(cells):
            continue
        pairs = []
        for index, value in enumerate(cells):
            header = headers[index] if index < len(headers) else f"column_{index + 1}"
            pairs.append(f"{header}={value}")
        lines.append(f"Row {row_index}: " + "; ".join(pairs))
    return "\n".join(lines) if lines else text.strip()


def _parse_table_rows(text: str) -> list[list[str]]:
    lines = [line.strip() for line in str(text or "").splitlines() if line.strip()]
    if not lines:
        return []
    if any("|" in line for line in lines):
        rows = []
        for line in lines:
            cleaned = line.strip("| ")
            cells = [cell.strip() for cell in cleaned.split("|")]
            if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells if cell):
                continue
            rows.append(cells)
        return rows
    delimiter = "\t" if any("\t" in line for line in lines) else ","
    try:
        return [[cell.strip() for cell in row] for row in csv.reader(io.StringIO("\n".join(lines)), delimiter=delimiter)]
    except Exception:
        return []


def _split_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        try:
            splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
        except Exception:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=max(chunk_size * 4, 1),
                chunk_overlap=max(chunk_overlap * 4, 0),
            )
        return [chunk.strip() for chunk in splitter.split_text(text) if chunk.strip()]
    except ImportError:
        return _fallback_split_text(text, max(chunk_size * 4, 1), max(chunk_overlap * 4, 0))


def _fallback_split_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not cleaned:
        return []
    chunks: list[str] = []
    start = 0
    overlap = min(chunk_overlap, max(chunk_size - 1, 0))
    while start < len(cleaned):
        end = min(start + chunk_size, len(cleaned))
        chunks.append(cleaned[start:end])
        if end >= len(cleaned):
            break
        start = max(end - overlap, start + 1)
    return chunks


def _to_int(value: Any, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default
