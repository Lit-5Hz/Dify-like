from __future__ import annotations

from pathlib import Path
from typing import Any
from uuid import uuid4

from fastapi import UploadFile
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.db.models import App, KnowledgeBase, KnowledgeChunk, KnowledgeDocument
from app.db.session import SessionLocal
from app.services.model_credential_service import get_model_credential
from app.services.qdrant_service import ensure_collection, search_knowledge_chunks, upsert_knowledge_chunks
from app.services.rag_providers import (
    build_embedding_provider,
    normalize_provider,
    rerank_chunks,
    validate_embedding_dimensions,
)


DIRECT_TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".py",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".java",
    ".go",
    ".json",
    ".yaml",
    ".yml",
    ".csv",
    ".html",
    ".css",
}
SUPPORTED_SUFFIXES = DIRECT_TEXT_SUFFIXES | {".pdf", ".docx"}
CODE_SUFFIXES = {".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".go", ".css", ".html"}


def list_runtime_knowledge_documents(
    db: Session,
    app_id: str,
    conversation_id: str,
    owner_user_id: str,
) -> list[KnowledgeDocument]:
    kb = get_runtime_knowledge_base(db, app_id, conversation_id, owner_user_id)
    if not kb:
        return []
    return list(
        db.scalars(
            select(KnowledgeDocument)
            .where(KnowledgeDocument.knowledge_base_id == kb.id)
            .order_by(KnowledgeDocument.created_at.desc())
        )
    )


def get_runtime_knowledge_base(
    db: Session,
    app_id: str,
    conversation_id: str,
    owner_user_id: str,
) -> KnowledgeBase | None:
    if not app_id or not conversation_id:
        return None
    return db.scalar(
        select(KnowledgeBase).where(
            KnowledgeBase.owner_user_id == owner_user_id,
            KnowledgeBase.scope == "runtime",
            KnowledgeBase.app_id == app_id,
            KnowledgeBase.conversation_id == conversation_id,
        )
    )


def get_or_create_runtime_knowledge_base(
    db: Session,
    app: App,
    conversation_id: str,
    rag_node: dict[str, Any],
) -> KnowledgeBase:
    existing = get_runtime_knowledge_base(db, app.id, conversation_id, app.owner_user_id)
    if existing:
        return existing

    config = _resolve_runtime_embedding_config(db, app.owner_user_id, rag_node)
    kb = KnowledgeBase(
        owner_user_id=app.owner_user_id,
        scope="runtime",
        app_id=app.id,
        conversation_id=conversation_id,
        name=f"{app.name} runtime files",
        description="Runtime files uploaded by the app user for this conversation.",
        embedding_provider=config["embedding_provider"],
        embedding_model=config["embedding_model"],
        embedding_dimension=config["embedding_dimension"],
        embedding_credential_id=config["embedding_credential_id"],
        embedding_base_url=config["embedding_base_url"],
        qdrant_collection=f"kb_{uuid4().hex}",
        chunk_size=config["chunk_size"],
        chunk_overlap=config["chunk_overlap"],
        chunk_strategy=config["chunk_strategy"],
        enable_parent_child=False,
        config_json={"runtime": True},
    )
    db.add(kb)
    db.flush()
    ensure_collection(kb)
    db.commit()
    db.refresh(kb)
    return kb


def rag_node_allows_user_upload(rag_node: dict[str, Any]) -> bool:
    return bool(rag_node.get("enabled", True))


async def save_knowledge_document(db: Session, kb: KnowledgeBase, file: UploadFile) -> KnowledgeDocument:
    filename = Path(file.filename or "").name
    suffix = Path(filename).suffix.lower()
    if not filename:
        raise ValueError("Filename is required.")
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError("Only txt/md/code/pdf/docx files are supported in the first RAG version.")

    settings = get_settings()
    kb_dir = Path(settings.storage_dir) / "knowledge_bases" / kb.id
    kb_dir.mkdir(parents=True, exist_ok=True)
    raw = await file.read()
    file_path = kb_dir / f"{uuid4()}_{filename}"
    file_path.write_bytes(raw)

    document = KnowledgeDocument(
        knowledge_base_id=kb.id,
        filename=filename,
        file_path=str(file_path),
        mime_type=file.content_type or "",
        status="queued",
        metadata_json={"suffix": suffix},
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    enqueue_knowledge_document(document.id)
    db.refresh(document)
    return document


def enqueue_knowledge_document(document_id: str) -> None:
    if get_settings().rag_inline_ingest:
        process_knowledge_document_sync(document_id)
        return
    from app.worker.celery_app import process_knowledge_document

    process_knowledge_document.delay(document_id)


def process_knowledge_document_sync(document_id: str) -> None:
    db = SessionLocal()
    try:
        _process_knowledge_document(db, document_id)
    finally:
        db.close()


def retrieve_rag_chunks(
    db: Session,
    app: App,
    owner_user_id: str,
    conversation_id: str,
    query: str,
    rag_node: dict[str, Any],
) -> dict[str, Any]:
    enabled = bool(rag_node.get("enabled", True))
    metadata: dict[str, Any] = {
        "kb_ids": [],
        "runtime_kb_id": None,
        "retrieval_mode": "disabled",
        "total_retrieved": 0,
        "total_returned": 0,
        "intent_matched": False,
        "standard_query": None,
        "warnings": [],
    }
    if not enabled:
        return {"chunks": [], "metadata": metadata}

    runtime_kb = get_runtime_knowledge_base(db, app.id, conversation_id, app.owner_user_id)
    if not runtime_kb:
        metadata["retrieval_mode"] = "empty"
        metadata["warnings"] = ["No files have been uploaded to this conversation RAG node yet."]
        return {"chunks": [], "metadata": metadata}

    chunks: list[dict[str, Any]] = []
    ensure_collection(runtime_kb)
    embedding_provider = build_embedding_provider(db, runtime_kb)
    query_vector = embedding_provider.embed_query(query)
    validate_embedding_dimensions([query_vector], runtime_kb.embedding_dimension)
    raw_hits = search_knowledge_chunks(runtime_kb, query_vector, max(int(rag_node.get("retrieval_top_k", 20)), 0))
    chunks.extend(_hit_to_chunk(hit, runtime_kb) for hit in raw_hits)

    ranked, warnings = rerank_chunks(
        db=db,
        owner_user_id=owner_user_id,
        query=query,
        chunks=chunks,
        provider=str(rag_node.get("rerank_provider") or "passthrough"),
        top_n=max(int(rag_node.get("rerank_top_n", 5)), 0),
        credential_id=str(rag_node.get("rerank_credential_id") or ""),
        model=str(rag_node.get("rerank_model") or ""),
        base_url=str(rag_node.get("rerank_base_url") or ""),
    )
    mode = "dense" if normalize_provider(str(rag_node.get("rerank_provider") or "passthrough")) in {"", "none", "passthrough"} else "dense+rerank"
    if warnings:
        mode = "dense+passthrough"

    metadata.update(
        {
            "kb_ids": [runtime_kb.id],
            "runtime_kb_id": runtime_kb.id,
            "retrieval_mode": mode,
            "total_retrieved": len(chunks),
            "total_returned": len(ranked),
            "warnings": warnings,
        }
    )
    return {"chunks": ranked, "metadata": metadata}


def _resolve_runtime_embedding_config(db: Session, owner_user_id: str, rag_node: dict[str, Any]) -> dict[str, Any]:
    provider = normalize_provider(str(rag_node.get("embedding_provider") or ""))
    model = str(rag_node.get("embedding_model") or "").strip()
    dimension = _to_int(rag_node.get("embedding_dimension"), 0)
    credential_id = str(rag_node.get("embedding_credential_id") or "").strip()
    base_url = str(rag_node.get("embedding_base_url") or "").strip()
    chunk_size = _to_int(rag_node.get("chunk_size"), 512)
    chunk_overlap = _to_int(rag_node.get("chunk_overlap"), 64)
    chunk_strategy = str(rag_node.get("chunk_strategy") or "auto").strip() or "auto"

    if not provider:
        raise ValueError("RAG node embedding provider is required before runtime upload.")
    if not model:
        raise ValueError("RAG node embedding model is required before runtime upload.")
    if dimension <= 0:
        raise ValueError("RAG node embedding dimension is required before runtime upload.")
    if chunk_overlap >= chunk_size:
        raise ValueError("RAG node chunk_overlap must be smaller than chunk_size.")
    if not get_model_credential(db, credential_id, owner_user_id):
        raise ValueError("RAG node embedding credential not found.")

    return {
        "embedding_provider": provider,
        "embedding_model": model,
        "embedding_dimension": dimension,
        "embedding_credential_id": credential_id,
        "embedding_base_url": base_url,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "chunk_strategy": chunk_strategy,
    }


def _to_int(value: Any, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _process_knowledge_document(db: Session, document_id: str) -> None:
    document = db.get(KnowledgeDocument, document_id)
    if not document:
        return
    kb = document.knowledge_base
    try:
        _update_document_status(db, document, "parsing")
        elements = _parse_document(Path(document.file_path), document.filename)
        _update_document_status(db, document, "chunking")
        chunks = _store_chunks(db, kb, document, elements)
        _update_document_status(db, document, "embedding")
        provider = build_embedding_provider(db, kb)
        vectors = provider.embed_documents([chunk.content for chunk in chunks])
        validate_embedding_dimensions(vectors, kb.embedding_dimension)
        for chunk in chunks:
            chunk.qdrant_point_id = chunk.id
        db.flush()
        ensure_collection(kb)
        upsert_knowledge_chunks(kb, chunks, vectors)
        document.status = "ready"
        document.error = ""
        kb.locked = True
        db.commit()
    except Exception as exc:
        document.status = "failed"
        document.error = str(exc)
        db.commit()


def _update_document_status(db: Session, document: KnowledgeDocument, status: str) -> None:
    document.status = status
    document.error = ""
    db.commit()
    db.refresh(document)


def _parse_document(path: Path, filename: str) -> list[dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix in DIRECT_TEXT_SUFFIXES:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [{"text": text, "chunk_type": "code" if suffix in CODE_SUFFIXES else "text", "section": None, "page_num": None}]
    if suffix == ".pdf":
        try:
            from pdfminer.high_level import extract_text
        except ImportError as exc:
            raise RuntimeError("pdfminer.six is not installed. Run `pip install -e .` in backend first.") from exc
        text = extract_text(str(path)) or ""
        return [{"text": text, "chunk_type": "text", "section": None, "page_num": None}]
    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError("python-docx is not installed. Run `pip install -e .` in backend first.") from exc
        doc = DocxDocument(str(path))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return [{"text": text, "chunk_type": "text", "section": None, "page_num": None}]
    raise ValueError(f"Unsupported file type for {filename}.")


def _store_chunks(
    db: Session,
    kb: KnowledgeBase,
    document: KnowledgeDocument,
    elements: list[dict[str, Any]],
) -> list[KnowledgeChunk]:
    db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_id == document.id))
    chunks: list[KnowledgeChunk] = []
    index = 0
    for element in elements:
        text = str(element.get("text") or "").strip()
        if not text:
            continue
        for content in _split_text(text, kb.chunk_size, kb.chunk_overlap):
            metadata = {
                "source_file": document.filename,
                "chunk_id": "",
                "parent_id": None,
                "chunk_type": element.get("chunk_type") or "text",
                "section": element.get("section"),
                "page_num": element.get("page_num"),
                "kb_id": kb.id,
                "scope": kb.scope,
                "app_id": kb.app_id,
                "conversation_id": kb.conversation_id,
            }
            chunk = KnowledgeChunk(
                knowledge_base_id=kb.id,
                document_id=document.id,
                chunk_index=index,
                content=content,
                metadata_json=metadata,
            )
            db.add(chunk)
            chunks.append(chunk)
            index += 1

    if not chunks:
        raise ValueError("No retrievable text was parsed from this document.")

    db.flush()
    for chunk in chunks:
        metadata = dict(chunk.metadata_json or {})
        metadata["chunk_id"] = chunk.id
        chunk.metadata_json = metadata
    db.commit()
    return chunks


def _split_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        try:
            splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
        except Exception:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=max(chunk_size * 4, 1),
                chunk_overlap=max(chunk_overlap * 4, 0),
            )
        return [chunk.strip() for chunk in splitter.split_text(text) if chunk.strip()]
    except ImportError:
        return _fallback_split_text(text, max(chunk_size * 4, 1), max(chunk_overlap * 4, 0))


def _fallback_split_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not cleaned:
        return []
    chunks: list[str] = []
    start = 0
    overlap = min(chunk_overlap, max(chunk_size - 1, 0))
    while start < len(cleaned):
        end = min(start + chunk_size, len(cleaned))
        chunks.append(cleaned[start:end])
        if end >= len(cleaned):
            break
        start = max(end - overlap, start + 1)
    return chunks


def _hit_to_chunk(hit: dict[str, Any], kb: KnowledgeBase | None = None) -> dict[str, Any]:
    payload = hit.get("payload") or {}
    return {
        "content": payload.get("content", ""),
        "score": float(hit.get("score", 0.0) or 0.0),
        "source_file": payload.get("source_file", ""),
        "page_num": payload.get("page_num"),
        "chunk_type": payload.get("chunk_type", "text"),
        "chunk_id": payload.get("chunk_id") or hit.get("id", ""),
        "section": payload.get("section"),
        "kb_id": payload.get("kb_id") or (kb.id if kb else ""),
        "scope": payload.get("scope") or (kb.scope if kb else ""),
    }
